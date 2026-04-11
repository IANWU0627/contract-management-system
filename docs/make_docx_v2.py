#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合同管理系统 - 精美 Word 文档生成脚本 v2
对标企业级交付标准，支持完整内容转换
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 颜色主题 ──────────────────────────────────────────────────
THEME = {
    'primary':     RGBColor(0x1A, 0x56, 0xB0),   # 深蓝（标题）
    'primary_lt':  RGBColor(0xDB, 0xEA, 0xFF),   # 浅蓝（表头背景）
    'accent':      RGBColor(0x0E, 0x3A, 0x70),   # 深藏蓝（封面背景）
    'heading2_bg': RGBColor(0xE8, 0xF1, 0xFF),   # H2 背景
    'table_even':  RGBColor(0xF5, 0xF8, 0xFF),   # 表格偶数行
    'code_bg':     RGBColor(0xF4, 0xF4, 0xF4),   # 代码块背景
    'border':      RGBColor(0xB0, 0xC8, 0xE8),   # 表格边框
    'text':        RGBColor(0x1E, 0x1E, 0x1E),   # 正文文字
    'text_sub':    RGBColor(0x55, 0x55, 0x55),   # 次要文字
    'white':       RGBColor(0xFF, 0xFF, 0xFF),
    'gold':        RGBColor(0xD4, 0xAF, 0x37),   # 金色（封面装饰）
    'green_bg':    RGBColor(0xE8, 0xF8, 0xEC),   # 绿色提示背景
    'orange_bg':   RGBColor(0xFE, 0xF3, 0xE2),   # 橙色提示背景
}

FONT_MAIN = '微软雅黑'
FONT_MONO = 'Courier New'


def hex_to_rgb(hex_str):
    h = hex_str.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def rgb_hex(rgb: RGBColor) -> str:
    """将 RGBColor 转为 6 位十六进制字符串（大写）"""
    # python-docx RGBColor.__str__ 直接返回类似 '1A56B0' 的十六进制字符串
    return str(rgb).upper()


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), rgb_hex(color))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_para_borders(para, top_color: RGBColor = None, left_color: RGBColor = None, left_sz=24):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    if top_color:
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), rgb_hex(top_color))
        pBdr.append(top)
    if left_color:
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), str(left_sz))
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), rgb_hex(left_color))
        pBdr.append(left)
    pPr.append(pBdr)


def set_para_shading(para, fill: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(fill))
    pPr.append(shd)


def add_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    run.font.color.rgb = THEME['text_sub']
    run.font.size = Pt(9)
    run.font.name = FONT_MAIN


def setup_document(doc, title_text):
    """设置文档基础样式"""
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 正文样式
    style = doc.styles['Normal']
    style.font.name = FONT_MAIN
    style.font.size = Pt(10.5)
    style.font.color.rgb = THEME['text']
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)

    add_page_number(doc)


def add_cover(doc, title, subtitle, meta_rows):
    """添加精美封面"""
    # 封面深色背景段落
    cover_bg = doc.add_paragraph()
    cover_bg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(cover_bg, THEME['accent'])
    cover_bg.paragraph_format.space_before = Pt(0)
    cover_bg.paragraph_format.space_after = Pt(0)

    # 顶部装饰条（金色）
    deco = doc.add_paragraph()
    set_para_shading(deco, THEME['gold'])
    deco.paragraph_format.space_before = Pt(0)
    deco.paragraph_format.space_after = Pt(0)
    run = deco.add_run(' ')
    run.font.size = Pt(6)

    # 封面背景块 - 产品名
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(brand, THEME['accent'])
    brand.paragraph_format.space_before = Pt(60)
    brand.paragraph_format.space_after = Pt(0)
    run = brand.add_run('Toy Contract')
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.color.rgb = THEME['gold']
    run.font.bold = True

    # 主标题
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(main_title, THEME['accent'])
    main_title.paragraph_format.space_before = Pt(20)
    main_title.paragraph_format.space_after = Pt(0)
    run = main_title.add_run(title)
    run.font.name = FONT_MAIN
    run.font.size = Pt(28)
    run.font.color.rgb = THEME['white']
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)

    # 副标题
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(sub, THEME['accent'])
    sub.paragraph_format.space_before = Pt(10)
    sub.paragraph_format.space_after = Pt(0)
    run = sub.add_run(subtitle)
    run.font.name = FONT_MAIN
    run.font.size = Pt(14)
    run.font.color.rgb = THEME['gold']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)

    # 分隔线
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(sep, THEME['accent'])
    sep.paragraph_format.space_before = Pt(20)
    sep.paragraph_format.space_after = Pt(20)
    run = sep.add_run('─' * 30)
    run.font.color.rgb = THEME['border']
    run.font.size = Pt(10)

    # 元数据表格（封面信息）
    tbl = doc.add_table(rows=len(meta_rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta_rows):
        row = tbl.rows[i]
        # 背景
        for cell in row.cells:
            set_cell_bg(cell, THEME['accent'])
        # key
        kc = row.cells[0]
        kc.width = Cm(4)
        kp = kc.paragraphs[0]
        kp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        kr = kp.add_run(k)
        kr.font.name = FONT_MAIN
        kr.font.size = Pt(11)
        kr.font.color.rgb = THEME['gold']
        kr.font.bold = True
        kr._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)
        # value
        vc = row.cells[1]
        vc.width = Cm(8)
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run('  ' + v)
        vr.font.name = FONT_MAIN
        vr.font.size = Pt(11)
        vr.font.color.rgb = THEME['white']
        vr._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)

    # 底部装饰
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(foot, THEME['accent'])
    foot.paragraph_format.space_before = Pt(40)
    foot.paragraph_format.space_after = Pt(0)
    run = foot.add_run('ContractHub  ·  2024')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = THEME['text_sub']

    # 金色底部线
    deco2 = doc.add_paragraph()
    set_para_shading(deco2, THEME['gold'])
    deco2.paragraph_format.space_before = Pt(20)
    deco2.paragraph_format.space_after = Pt(0)
    run = deco2.add_run(' ')
    run.font.size = Pt(6)

    # 分节符
    doc.add_page_break()


def add_heading1(doc, text):
    """一级标题：左侧蓝色竖条 + 蓝色文字"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(8)
    set_para_shading(para, THEME['primary_lt'])
    set_para_borders(para, left_color=THEME['primary'], left_sz=28)
    run = para.add_run('  ' + text)
    run.font.name = FONT_MAIN
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = THEME['primary']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)
    return para


def add_heading2(doc, text):
    """二级标题：蓝色菱形 + 深色文字"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    set_para_borders(para, top_color=THEME['border'])
    run = para.add_run('◆ ' + text)
    run.font.name = FONT_MAIN
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = THEME['primary']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)
    return para


def add_heading3(doc, text):
    """三级标题"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run('▸ ' + text)
    run.font.name = FONT_MAIN
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = THEME['primary']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)
    return para


def add_heading4(doc, text):
    """四级标题"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run('• ' + text)
    run.font.name = FONT_MAIN
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = THEME['text']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)
    return para


def add_body(doc, text, indent=0):
    """正文段落"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = Pt(18)
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    run = para.add_run(text)
    run.font.name = FONT_MAIN
    run.font.size = Pt(10.5)
    run.font.color.rgb = THEME['text']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)
    return para


def add_bullet(doc, text, level=0):
    """列表项"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    bullets = ['●', '○', '▪']
    bullet_char = bullets[min(level, 2)]
    run = para.add_run(f'{bullet_char}  {text}')
    run.font.name = FONT_MAIN
    run.font.size = Pt(10.5)
    run.font.color.rgb = THEME['text']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)
    return para


def add_code_block(doc, text):
    """代码块：灰色背景 + 等宽字体"""
    # 先加一个小空行
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)
    sp.paragraph_format.space_before = Pt(4)

    for line in text.split('\n'):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.right_indent = Cm(0.5)
        set_para_shading(para, THEME['code_bg'])
        run = para.add_run(line if line else ' ')
        run.font.name = FONT_MONO
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_after = Pt(4)
    sp2.paragraph_format.space_before = Pt(0)


def add_info_box(doc, text, box_type='info'):
    """彩色信息框"""
    bg_map = {
        'info': THEME['primary_lt'],
        'warning': THEME['orange_bg'],
        'success': THEME['green_bg'],
        'code': THEME['code_bg'],
    }
    border_map = {
        'info': THEME['primary'],
        'warning': RGBColor(0xE6, 0xA2, 0x3C),
        'success': RGBColor(0x67, 0xC2, 0x3A),
        'code': THEME['text_sub'],
    }
    bg = bg_map.get(box_type, THEME['primary_lt'])
    border = border_map.get(box_type, THEME['primary'])

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Cm(0.3)
    set_para_shading(para, bg)
    set_para_borders(para, left_color=border, left_sz=20)
    run = para.add_run('  ' + text)
    run.font.name = FONT_MAIN
    run.font.size = Pt(10)
    run.font.color.rgb = THEME['text']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)
    return para


def add_table(doc, headers, rows, col_widths=None):
    """添加样式化表格"""
    n_cols = len(headers)
    n_rows = len(rows)
    tbl = doc.add_table(rows=1 + n_rows, cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    # 表头行
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, THEME['primary'])
        set_cell_borders(cell,
            top=THEME['primary'], bottom=THEME['primary'],
            left=THEME['white'], right=THEME['white'])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(str(h))
        run.font.name = FONT_MAIN
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = THEME['white']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg = THEME['table_even'] if r_idx % 2 == 0 else THEME['white']
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_borders(cell,
                top=THEME['border'], bottom=THEME['border'],
                left=THEME['border'], right=THEME['border'])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(str(cell_text) if cell_text else '')
            run.font.name = FONT_MAIN
            run.font.size = Pt(10)
            run.font.color.rgb = THEME['text']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)

    # 表格后空行
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


def parse_md_table(lines):
    """解析 Markdown 表格，返回 (headers, rows)"""
    headers = []
    rows = []
    for i, line in enumerate(lines):
        line = line.strip()
        if not line.startswith('|'):
            continue
        parts = [p.strip() for p in line.strip('|').split('|')]
        if i == 0:
            headers = parts
        elif re.match(r'^[-:\s|]+$', line):
            continue
        else:
            rows.append(parts)
    return headers, rows


def process_markdown_to_doc(doc, md_text):
    """将 Markdown 文本转换为 Word 内容"""
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        raw = line

        # 代码块
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
                i += 1
                continue
            else:
                add_code_block(doc, '\n'.join(code_lines))
                in_code = False
                code_lines = []
                i += 1
                continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # 表格
        stripped = line.strip()
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines.append(line)
            i += 1
            # 收集完整表格
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_md_table(table_lines)
            if headers:
                add_table(doc, headers, rows)
            table_lines = []
            continue

        # 标题
        if stripped.startswith('#### '):
            add_heading4(doc, stripped[5:].strip())
        elif stripped.startswith('### '):
            add_heading3(doc, stripped[4:].strip())
        elif stripped.startswith('## '):
            add_heading2(doc, stripped[3:].strip())
        elif stripped.startswith('# '):
            # 文档内部一级标题（跳过，封面已处理）
            pass
        # 水平线
        elif stripped.startswith('---'):
            sp = doc.add_paragraph()
            set_para_borders(sp, top_color=THEME['border'])
            sp.paragraph_format.space_before = Pt(8)
            sp.paragraph_format.space_after = Pt(8)
        # 列表
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            # 处理加粗
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'[\1]', text)
            indent = (len(line) - len(line.lstrip())) // 2
            add_bullet(doc, text, level=indent)
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            add_bullet(doc, text, level=0)
        # 加粗段落（如 **xxx**:）
        elif stripped.startswith('**') and ('**:' in stripped or '**：' in stripped):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(2)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            run = para.add_run(text)
            run.font.name = FONT_MAIN
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = THEME['primary']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_MAIN)
        # 引用块
        elif stripped.startswith('> '):
            add_info_box(doc, stripped[2:].strip(), 'info')
        # 空行
        elif not stripped:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after = Pt(2)
        # 普通段落
        else:
            # 清除 markdown 标记
            text = stripped
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'[\1]', text)
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
            if text:
                add_body(doc, text)

        i += 1


def generate_doc(md_file, out_file, cover_title, cover_subtitle, cover_meta):
    """生成单个 Word 文档"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 跳过文档第一行（H1 标题，已在封面显示）
    lines = content.split('\n')
    body_start = 0
    for j, ln in enumerate(lines):
        if ln.strip().startswith('# '):
            body_start = j + 1
            break
    body_content = '\n'.join(lines[body_start:])

    doc = Document()
    setup_document(doc, cover_title)
    add_cover(doc, cover_title, cover_subtitle, cover_meta)
    process_markdown_to_doc(doc, body_content)
    doc.save(out_file)
    print(f'  ✅ 已生成：{os.path.basename(out_file)}  ({os.path.getsize(out_file) // 1024} KB)')


# ── 主程序 ────────────────────────────────────────────────────
if __name__ == '__main__':
    docs_dir = os.path.dirname(os.path.abspath(__file__))
    print('\n🚀 开始生成企业级交付文档（v2）...\n')

    documents = [
        {
            'md': '01-技术选型文档.md',
            'out': '01-技术选型文档.docx',
            'title': '技术选型文档',
            'subtitle': 'Technology Selection Document',
            'meta': [
                ('文档版本', 'v2.0.0'),
                ('项目名称', 'Toy Contract 合同管理系统'),
                ('文档状态', '已批准'),
                ('编写团队', 'ContractHub 架构组'),
                ('最后更新', '2024年12月'),
            ]
        },
        {
            'md': '02-模块设计文档.md',
            'out': '02-模块设计文档.docx',
            'title': '模块设计文档',
            'subtitle': 'Module Design Document',
            'meta': [
                ('文档版本', 'v2.0.0'),
                ('项目名称', 'Toy Contract 合同管理系统'),
                ('文档状态', '已批准'),
                ('编写团队', 'ContractHub 架构组'),
                ('最后更新', '2024年12月'),
            ]
        },
        {
            'md': '03-开发路线图.md',
            'out': '03-开发路线图.docx',
            'title': '开发路线图',
            'subtitle': 'Development Roadmap',
            'meta': [
                ('文档版本', 'v2.0.0'),
                ('项目名称', 'Toy Contract 合同管理系统'),
                ('项目周期', '12 周（2024-03 至 2024-05）'),
                ('编写团队', 'ContractHub 项目组'),
                ('最后更新', '2024年12月'),
            ]
        },
        {
            'md': '04-BRD业务需求文档.md',
            'out': '04-BRD业务需求文档.docx',
            'title': 'BRD 业务需求文档',
            'subtitle': 'Business Requirements Document',
            'meta': [
                ('文档版本', 'v2.0.0'),
                ('项目名称', 'Toy Contract 合同管理系统'),
                ('文档状态', '已批准'),
                ('目标读者', '业务方 / 产品 / 开发 / 投资方'),
                ('最后更新', '2024年12月'),
            ]
        },
        {
            'md': '05-PRD产品需求文档.md',
            'out': '05-PRD产品需求文档.docx',
            'title': 'PRD 产品需求文档',
            'subtitle': 'Product Requirements Document',
            'meta': [
                ('文档版本', 'v2.0.0'),
                ('项目名称', 'Toy Contract 合同管理系统'),
                ('文档状态', '已批准'),
                ('目标读者', '设计 / 前端 / 后端 / 测试'),
                ('最后更新', '2024年12月'),
            ]
        },
    ]

    for d in documents:
        md_path = os.path.join(docs_dir, d['md'])
        out_path = os.path.join(docs_dir, d['out'])
        generate_doc(md_path, out_path, d['title'], d['subtitle'], d['meta'])

    print('\n🎉 全部文档生成完成！\n')
