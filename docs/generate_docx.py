#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成项目文档Word文件
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def add_paragraph_custom(doc, text, bold=False, size=11):
    """添加自定义段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def create_feature_table(doc, headers, rows):
    """创建功能表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], '4472C4')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    return table

def generate_document():
    """生成项目文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Toy Contract\n企业合同管理系统')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('\n项目功能说明书\n')
    run.font.size = Pt(24)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 日期
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('\n\n2026年3月')
    run.font.size = Pt(14)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # 目录
    add_heading_custom(doc, '目录', level=1)
    toc_items = [
        '1. 项目概述',
        '2. 功能模块',
        '3. 技术架构',
        '4. 用户角色与权限',
        '5. 数据库设计',
        '6. API接口规范',
        '7. 安全特性',
        '8. 部署说明'
    ]
    for item in toc_items:
        add_paragraph_custom(doc, item)
    
    doc.add_page_break()
    
    # 1. 项目概述
    add_heading_custom(doc, '1. 项目概述', level=1)
    
    add_heading_custom(doc, '1.1 产品定位', level=2)
    add_paragraph_custom(doc, 'Toy Contract 是一款面向企业的合同全生命周期管理平台，提供从合同创建、审批、签署到归档的一站式解决方案。')
    
    add_heading_custom(doc, '1.2 目标用户', level=2)
    users = ['企业法务部门', '采购/销售部门', '行政管理人员', '企业高管']
    for user in users:
        add_paragraph_custom(doc, f'• {user}')
    
    add_heading_custom(doc, '1.3 核心价值', level=2)
    values = [
        '提升合同管理效率 - 数字化管理，减少纸质文档',
        '降低合同风险 - 审批流程规范化，风险可控',
        '规范审批流程 - 自定义审批节点，流程透明',
        '数据可视化分析 - 多维度统计报表，决策支持'
    ]
    for value in values:
        add_paragraph_custom(doc, f'• {value}')
    
    doc.add_page_break()
    
    # 2. 功能模块
    add_heading_custom(doc, '2. 功能模块', level=1)
    
    # 2.1 用户管理
    add_heading_custom(doc, '2.1 用户管理模块', level=2)
    add_paragraph_custom(doc, '提供完整的用户生命周期管理功能，支持多种角色权限控制。')
    
    headers = ['功能', '说明', '状态']
    rows = [
        ['用户注册/登录', 'JWT Token认证，支持记住密码', '✅ 已完成'],
        ['JWT Token刷新', '自动刷新机制，无感续期', '✅ 已完成'],
        ['角色权限控制', 'RBAC权限模型，细粒度控制', '✅ 已完成'],
        ['用户信息管理', '头像、昵称、联系方式管理', '✅ 已完成'],
        ['密码修改', 'BCrypt加密，安全存储', '✅ 已完成'],
        ['密码重置', '管理员可重置用户密码', '✅ 已完成'],
        ['用户状态管理', '启用/禁用用户账号', '✅ 已完成'],
        ['用户审计', '用户操作日志，登录历史记录', '✅ 已完成'],
        ['角色管理', '创建、编辑、删除角色，为角色分配权限', '✅ 已完成'],
        ['权限管理', '细粒度权限配置，权限分组管理', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.2 合同管理
    add_heading_custom(doc, '2.2 合同管理模块', level=2)
    add_paragraph_custom(doc, '核心功能模块，支持合同全生命周期管理。')
    
    add_heading_custom(doc, '合同类型', level=3)
    contract_types = ['采购合同 (PURCHASE)', '销售合同 (SALES)', '租赁合同 (LEASE)', 
                      '雇佣合同 (EMPLOYMENT)', '服务合同 (SERVICE)', '其他合同 (OTHER)']
    for ct in contract_types:
        add_paragraph_custom(doc, f'• {ct}')
    
    add_heading_custom(doc, '合同状态流转', level=3)
    add_paragraph_custom(doc, '草稿 → 待审批 → 已批准 → 已签署 → 已归档 → 已终止')
    
    add_heading_custom(doc, '功能列表', level=3)
    headers = ['功能', '说明', '状态']
    rows = [
        ['合同创建', '支持模板选择，快速创建，多方合同支持', '✅ 已完成'],
        ['合同编辑', '富文本编辑器，支持格式化，实时保存', '✅ 已完成'],
        ['合同删除', '软删除机制，数据可恢复，删除权限控制', '✅ 已完成'],
        ['合同详情', '完整信息展示，附件预览，版本历史', '✅ 已完成'],
        ['合同列表', '分页、多条件筛选、关键词搜索、排序', '✅ 已完成'],
        ['合同审批流程', '提交审批，多级审批，审批意见，审批历史', '✅ 已完成'],
        ['合同签署', 'PDF生成，签署状态跟踪，签署人管理', '✅ 已完成'],
        ['合同归档', '归档/解档操作，归档原因记录', '✅ 已完成'],
        ['合同终止', '终止合同，记录终止原因，终止日期', '✅ 已完成'],
        ['PDF生成/下载', '专业PDF格式，水印支持，批量下载', '✅ 已完成'],
        ['AI合同分析', '智能分析合同风险，关键条款提取，合规检查', '✅ 已完成'],
        ['版本历史记录', '完整版本追溯，支持附件版本，变更对比', '✅ 已完成'],
        ['合同评论', '协作评论，@提及功能，评论通知', '✅ 已完成'],
        ['合同附件管理', '多附件上传，文件类型限制，附件预览，附件下载', '✅ 已完成'],
        ['合同草稿自动保存', '定时自动保存，防数据丢失，草稿恢复', '✅ 已完成'],
        ['合同编号生成', '自动生成合同编号，支持自定义规则', '✅ 已完成'],
        ['合同标签', '多标签管理，标签说明，标签筛选，标签统计', '✅ 已完成'],
        ['提醒规则管理', '创建、编辑、删除提醒规则，规则优先级设置', '✅ 已完成'],
        ['合同续约管理', '续约提醒，续约申请，续约审批', '✅ 已完成'],
        ['合同收藏', '收藏重要合同，快速访问，批量管理', '✅ 已完成'],
        ['文件夹管理', '合同文件夹创建、编辑、删除，文件分类管理', '✅ 已完成'],
        ['变更日志', '合同变更历史记录，变更内容对比', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # 2.3 模板管理
    add_heading_custom(doc, '2.3 模板管理模块', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['模板列表', '分类筛选，关键词搜索，分页展示，使用次数统计', '✅ 已完成'],
        ['模板创建/编辑/删除', '富文本编辑，模板分类，模板状态管理', '✅ 已完成'],
        ['模板变量定义', 'JSON格式定义，变量类型支持，变量默认值', '✅ 已完成'],
        ['变量替换功能', '{{变量名}}语法，实时替换预览，批量替换', '✅ 已完成'],
        ['实时预览', '变量填充效果预览，PDF预览，HTML预览', '✅ 已完成'],
        ['从模板创建合同', '快速创建，变量自动填充，模板版本选择', '✅ 已完成'],
        ['模板使用次数统计', '使用频率分析，热门模板推荐', '✅ 已完成'],
        ['模板复制', '快速复制现有模板，修改后保存为新模板', '✅ 已完成'],
        ['模板水印功能', 'PDF水印设置，水印位置调整，水印样式自定义', '✅ 已完成'],
        ['模板导出', '导出为Word，导出为PDF，导出模板变量', '✅ 已完成'],
        ['模板权限控制', '模板可见性设置，模板编辑权限', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.4 变量管理模块
    add_heading_custom(doc, '2.4 变量管理模块', level=2)
    add_paragraph_custom(doc, '独立的变量管理系统，支持可视化变量配置，模板引用更便捷。')
    
    add_heading_custom(doc, '变量分类', level=3)
    variable_categories = [
        '🌐 系统变量 (8个) - 合同基础信息',
        '🏢 相对方信息 (10个) - 甲乙双方信息',
        '📦 采购合同 (7个) - 采购专用变量',
        '🛠️ 服务合同 (5个) - 服务专用变量',
        '🏠 租赁合同 (5个) - 租赁专用变量',
        '👔 劳动合同 (5个) - 劳动专用变量',
        '✏️ 自定义变量 (无限制) - 用户自定义'
    ]
    for category in variable_categories:
        add_paragraph_custom(doc, f'• {category}')
    
    add_heading_custom(doc, '功能列表', level=3)
    headers = ['功能', '说明', '状态']
    rows = [
        ['变量列表', '分类筛选，类型筛选，关键词搜索', '✅ 已完成'],
        ['变量CRUD', '创建、编辑、删除变量', '✅ 已完成'],
        ['变量分类管理', '按分类组织变量', '✅ 已完成'],
        ['变量类型配置', 'text/number/date/textarea', '✅ 已完成'],
        ['默认值设置', '可选的默认值', '✅ 已完成'],
        ['必填标识', '标记必填变量', '✅ 已完成'],
        ['批量初始化', '一键创建40+预定义变量', '✅ 已完成'],
        ['变量引用', '模板编辑器快速插入变量', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.5 审批管理
    add_heading_custom(doc, '2.5 审批管理模块', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['待审批列表', '按时间排序，审批优先级，批量审批', '✅ 已完成'],
        ['审批操作', '审批通过，审批拒绝，审批转办，审批评论', '✅ 已完成'],
        ['审批历史', '完整审批记录，审批意见查看，审批时间线', '✅ 已完成'],
        ['审批通知', '审批任务提醒，审批结果通知，逾期提醒', '✅ 已完成'],
        ['审批流程配置', '自定义审批节点，审批人设置，审批条件设置', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.6 到期提醒
    add_heading_custom(doc, '2.6 到期提醒模块', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['定时任务自动检测', '每天9点自动执行，检测即将到期合同', '✅ 已完成'],
        ['到期检测规则', '30天、7天，当天到期提醒，提前提醒天数可配置', '✅ 已完成'],
        ['提醒列表', '按到期时间排序，提醒状态管理，批量处理', '✅ 已完成'],
        ['手动触发检查', '立即执行检查，手动刷新提醒列表', '✅ 已完成'],
        ['提醒规则管理', '规则名称，适用合同类型，金额范围，提醒天数，启用状态', '✅ 已完成'],
        ['提醒通知', '系统内通知，邮件通知（待实现），短信通知（待实现）', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.7 统计报表
    add_heading_custom(doc, '2.7 统计报表模块', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['合同总数统计', '按时间段，按部门，按类型统计', '✅ 已完成'],
        ['合同金额统计', '总金额，平均金额，最大/最小金额', '✅ 已完成'],
        ['按类型统计', '各类型合同数量，占比分析，趋势变化', '✅ 已完成'],
        ['按状态统计', '各状态合同数量，处理效率分析', '✅ 已完成'],
        ['趋势图表', 'ECharts可视化，折线图，饼图，柱状图', '✅ 已完成'],
        ['合同导出Excel', '自定义导出字段，批量导出，导出模板', '✅ 已完成'],
        ['报表筛选', '多条件组合筛选，时间范围选择，部门筛选', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # 2.8 标签管理
    add_heading_custom(doc, '2.8 标签管理模块', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['标签列表', '标签名称，标签颜色，标签说明，使用次数', '✅ 已完成'],
        ['标签创建/编辑/删除', '标签名称，颜色选择，说明填写', '✅ 已完成'],
        ['标签说明字段', '详细描述标签用途，使用场景说明', '✅ 已完成'],
        ['合同标签关联', '批量添加标签，标签筛选，标签统计', '✅ 已完成'],
        ['标签分类', '标签分组管理，分类统计，分类筛选', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.9 续约管理
    add_heading_custom(doc, '2.9 续约管理模块', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['续约列表', '待续约合同，已续约合同，续约状态管理', '✅ 已完成'],
        ['续约创建', '选择合同，设置新期限，续约类型，续约原因', '✅ 已完成'],
        ['续约审批', '续约申请审批，审批流程，审批历史', '✅ 已完成'],
        ['续约历史', '合同续约记录，续约详情，续约趋势', '✅ 已完成'],
        ['续约提醒', '即将到期合同提醒，续约截止日期提醒', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.10 收藏管理
    add_heading_custom(doc, '2.10 收藏管理模块', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['我的收藏列表', '收藏合同展示，收藏时间，收藏分类', '✅ 已完成'],
        ['添加收藏', '合同详情页收藏，批量收藏，收藏备注', '✅ 已完成'],
        ['取消收藏', '单个取消，批量取消，收藏管理', '✅ 已完成'],
        ['收藏分类', '自定义收藏夹，分类管理，分类统计', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.11 文件夹管理
    add_heading_custom(doc, '2.11 文件夹管理模块', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['文件夹创建/编辑/删除', '文件夹名称，父文件夹选择，权限设置', '✅ 已完成'],
        ['合同归类到文件夹', '拖拽归类，批量归类，自动归类规则', '✅ 已完成'],
        ['文件夹权限控制', '文件夹可见性，编辑权限，访问权限', '✅ 已完成'],
        ['文件夹树形结构', '层级展示，折叠/展开，路径导航', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.12 合同类型字段配置
    add_heading_custom(doc, '2.12 合同类型字段配置模块', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['自定义合同类型字段', '字段名称，字段类型，字段顺序', '✅ 已完成'],
        ['字段类型配置', '文本、数字、日期、下拉选择、多选、附件等', '✅ 已完成'],
        ['字段验证规则', '必填项，长度限制，格式验证，正则表达式', '✅ 已完成'],
        ['字段默认值', '静态默认值，动态默认值，条件默认值', '✅ 已完成'],
        ['字段权限', '字段可见性，编辑权限，必填性配置', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.13 操作日志
    add_heading_custom(doc, '2.13 操作日志模块', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['操作日志列表', '用户，时间，操作类型，操作模块，操作结果', '✅ 已完成'],
        ['日志筛选和搜索', '按用户，按时间，按操作类型，按模块', '✅ 已完成'],
        ['操作详情查看', '详细操作内容，操作前后对比，IP地址，浏览器信息', '✅ 已完成'],
        ['日志导出', '导出为Excel，导出为CSV，导出筛选结果', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 2.14 系统功能
    add_heading_custom(doc, '2.14 系统功能', level=2)
    headers = ['功能', '说明', '状态']
    rows = [
        ['操作日志（审计日志AOP）', '自动记录所有操作，不可篡改，安全存储', '✅ 已完成'],
        ['多语言切换', '中英文切换，语言包管理，自定义语言', '✅ 已完成'],
        ['主题切换', '亮色/暗色主题，主题自定义，主题保存', '✅ 已完成'],
        ['数据导出', 'Excel，PDF，CSV格式导出，导出模板管理', '✅ 已完成'],
        ['速率限制', 'API防护，60次/分钟，防止暴力攻击', '✅ 已完成'],
        ['CORS安全配置', '跨域安全配置，白名单管理', '✅ 已完成'],
        ['系统设置', '系统参数配置，邮件服务器配置，通知设置', '✅ 已完成'],
        ['数据备份', '手动备份，自动备份，备份恢复（待实现）', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # 3. 技术架构
    add_heading_custom(doc, '3. 技术架构', level=1)
    
    add_heading_custom(doc, '3.1 后端技术栈', level=2)
    headers = ['技术', '版本', '用途', '选型理由']
    rows = [
        ['Java', '21', '编程语言', '最新LTS版本，性能优化，语法特性丰富，生态成熟'],
        ['Spring Boot', '3.2.0', 'Web框架', '快速开发，生态完善，自动配置，内嵌服务器'],
        ['Spring Security', '6.2.0', '安全框架', '标准安全解决方案，支持JWT，权限控制'],
        ['MyBatis Plus', '3.5.5', 'ORM框架', '简化CRUD操作，代码生成，性能优化，易于使用'],
        ['H2 Database', '2.2.224', '开发数据库', '内存数据库，快速启动，适合开发测试，无需额外配置'],
        ['MySQL', '8.x', '生产数据库', '稳定可靠，性能优异，生态成熟，适合企业级应用'],
        ['JWT', '0.12.3', 'Token认证', '无状态认证，便于水平扩展，适合前后端分离架构'],
        ['iText', '8.0.2', 'PDF生成', '专业PDF处理库，功能强大，支持水印，易于集成'],
        ['Lombok', '1.18.30', '代码简化', '减少样板代码，提高开发效率，代码更简洁'],
        ['SLF4J/Logback', '-', '日志框架', '灵活配置，性能优异，支持多种输出方式'],
        ['AspectJ', '-', '切面编程', '用于审计日志，权限控制，事务管理'],
        ['Jackson', '-', 'JSON处理', '高性能JSON库，Spring Boot默认集成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.2 前端技术栈', level=2)
    headers = ['技术', '版本', '用途', '选型理由']
    rows = [
        ['Vue 3', '3.x', '前端框架', '响应式设计，组合式API，性能优异，生态成熟'],
        ['TypeScript', '5.x', '类型系统', '类型安全，代码提示，减少错误，提高代码质量'],
        ['Vite', '5.x', '构建工具', '快速启动，热更新，优化构建，开发体验好'],
        ['Element Plus', '2.x', 'UI组件库', '丰富的组件，美观的设计，易于使用，文档完善'],
        ['ECharts', '6.x', '图表库', '丰富的图表类型，交互性强，性能优异，文档完善'],
        ['Pinia', '3.x', '状态管理', '轻量级，TypeScript支持，易于集成，性能优异'],
        ['Vue Router', '5.x', '路由管理', '官方路由解决方案，支持动态路由，嵌套路由'],
        ['vue-i18n', '11.x', '国际化', '支持多语言，易于配置，文档完善'],
        ['Axios', '1.x', 'HTTP客户端', '功能强大，易于使用，拦截器支持，文档完善'],
        ['Day.js', '-', '日期处理', '轻量级，API友好，易于使用'],
        ['Lodash', '-', '工具库', '丰富的工具函数，提高开发效率'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.3 数据库索引设计', level=2)
    add_paragraph_custom(doc, '为提升查询性能，已在以下字段添加索引：')
    index_list = [
        'contract：contract_no (唯一索引), type, status, end_date, creator_id, create_time',
        'user：username (唯一索引), role, status',
        'contract_template：name, category',
        'contract_version：contract_id, created_at',
        'contract_reminder：contract_id, expire_date, status',
        'approval_record：contract_id, approver_id, create_time',
        'operation_log：user_id, action, module, create_time',
        'contract_tag：name, color',
        'contract_renewal：contract_id, status, new_end_date',
        'contract_folder：name, parent_id',
        'reminder_rule：name, contract_type, is_enabled',
    ]
    for index in index_list:
        add_paragraph_custom(doc, f'• {index}')
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.4 系统架构特点', level=2)
    features = [
        '前后端分离：前端Vue 3 + 后端Spring Boot，独立部署，易于扩展',
        '微服务架构：模块化设计，服务解耦，便于独立开发和部署',
        '安全可靠：JWT认证，RBAC权限控制，审计日志，API防护',
        '高性能：数据库索引优化，缓存策略，异步处理',
        '可扩展：模块化设计，插件机制，易于添加新功能',
    ]
    for feature in features:
        add_paragraph_custom(doc, f'• {feature}')
    
    doc.add_page_break()
    
    # 4. 用户角色与权限
    add_heading_custom(doc, '4. 用户角色与权限', level=1)
    
    add_heading_custom(doc, '4.1 角色定义', level=2)
    headers = ['角色', '权限范围', '核心功能']
    rows = [
        ['ADMIN (管理员)', '全部功能，包括用户管理、系统设置、变量管理', '用户管理、角色管理、系统配置、权限管理、变量管理'],
        ['LEGAL (法务)', '合同管理、审批、模板、统计、提醒、续约、变量查看', '合同审批、模板管理、统计报表、续约管理'],
        ['USER (普通用户)', '合同查看、创建（自己的合同）、模板使用、收藏', '合同创建、查看、收藏'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '4.2 权限清单', level=2)
    headers = ['权限名称', '权限代码', '说明']
    rows = [
        ['合同管理', 'CONTRACT_MANAGE', '管理合同的增删改查'],
        ['合同审批', 'CONTRACT_APPROVE', '审批合同'],
        ['模板管理', 'TEMPLATE_MANAGE', '管理合同模板'],
        ['用户管理', 'USER_MANAGE', '管理用户'],
        ['角色管理', 'ROLE_MANAGE', '管理角色'],
        ['权限管理', 'PERMISSION_MANAGE', '管理权限'],
        ['统计报表', 'REPORT_VIEW', '查看统计报表'],
        ['提醒管理', 'REMINDER_MANAGE', '管理合同到期提醒'],
        ['我的收藏', 'FAVORITE_MANAGE', '管理收藏的合同'],
        ['续约管理', 'RENEWAL_MANAGE', '管理合同续约'],
        ['标签管理', 'TAG_MANAGE', '管理合同标签'],
        ['提醒规则', 'REMINDER_RULE_MANAGE', '管理提醒规则'],
        ['变量管理', 'VARIABLE_MANAGE', '管理模板变量'],
        ['合同类型字段配置', 'TYPE_FIELD_CONFIG_MANAGE', '管理合同类型字段'],
        ['系统设置', 'SETTING_MANAGE', '管理系统设置'],
        ['文件夹管理', 'FOLDER_MANAGE', '管理合同文件夹'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # 5. 数据库设计
    add_heading_custom(doc, '5. 数据库设计', level=1)
    
    add_heading_custom(doc, '5.1 核心表结构', level=2)
    
    add_heading_custom(doc, 'user (用户表)', level=3)
    headers = ['字段', '类型', '说明', '约束']
    rows = [
        ['id', 'BIGINT', '主键', 'AUTO_INCREMENT'],
        ['username', 'VARCHAR(50)', '用户名', 'NOT NULL UNIQUE'],
        ['password', 'VARCHAR(255)', '密码(BCrypt)', 'NOT NULL'],
        ['nickname', 'VARCHAR(50)', '昵称', '-'],
        ['email', 'VARCHAR(100)', '邮箱', '-'],
        ['phone', 'VARCHAR(20)', '电话', '-'],
        ['avatar', 'TEXT', '头像', '-'],
        ['role', 'VARCHAR(20)', '角色', 'DEFAULT \'USER\''],
        ['status', 'INT', '状态', 'DEFAULT 1'],
        ['create_time', 'TIMESTAMP', '创建时间', 'DEFAULT CURRENT_TIMESTAMP'],
        ['update_time', 'TIMESTAMP', '更新时间', 'DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, 'role (角色表)', level=3)
    headers = ['字段', '类型', '说明', '约束']
    rows = [
        ['id', 'BIGINT', '主键', 'AUTO_INCREMENT'],
        ['name', 'VARCHAR(50)', '角色名称', 'NOT NULL'],
        ['code', 'VARCHAR(50)', '角色代码', 'NOT NULL UNIQUE'],
        ['description', 'TEXT', '角色描述', '-'],
        ['status', 'INT', '状态', 'DEFAULT 1'],
        ['created_at', 'TIMESTAMP', '创建时间', 'DEFAULT CURRENT_TIMESTAMP'],
        ['updated_at', 'TIMESTAMP', '更新时间', 'DEFAULT CURRENT_TIMESTAMP'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, 'permission (权限表)', level=3)
    headers = ['字段', '类型', '说明', '约束']
    rows = [
        ['id', 'BIGINT', '主键', 'AUTO_INCREMENT'],
        ['name', 'VARCHAR(50)', '权限名称', 'NOT NULL'],
        ['code', 'VARCHAR(50)', '权限代码', 'NOT NULL UNIQUE'],
        ['path', 'VARCHAR(255)', 'API路径', '-'],
        ['method', 'VARCHAR(10)', 'HTTP方法', '-'],
        ['description', 'TEXT', '权限描述', '-'],
        ['status', 'INT', '状态', 'DEFAULT 1'],
        ['created_at', 'TIMESTAMP', '创建时间', 'DEFAULT CURRENT_TIMESTAMP'],
        ['updated_at', 'TIMESTAMP', '更新时间', 'DEFAULT CURRENT_TIMESTAMP'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, 'role_permission (角色权限关联表)', level=3)
    headers = ['字段', '类型', '说明', '约束']
    rows = [
        ['id', 'BIGINT', '主键', 'AUTO_INCREMENT'],
        ['role_id', 'BIGINT', '角色ID', 'NOT NULL'],
        ['permission_id', 'BIGINT', '权限ID', 'NOT NULL'],
        ['created_at', 'TIMESTAMP', '创建时间', 'DEFAULT CURRENT_TIMESTAMP'],
        ['UNIQUE KEY', '-', '唯一约束', '(role_id, permission_id)'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, 'contract (合同表)', level=3)
    headers = ['字段', '类型', '说明', '约束']
    rows = [
        ['id', 'BIGINT', '主键', 'AUTO_INCREMENT'],
        ['contract_no', 'VARCHAR(50)', '合同编号', 'NOT NULL UNIQUE'],
        ['title', 'VARCHAR(255)', '标题', 'NOT NULL'],
        ['type', 'VARCHAR(20)', '类型', 'NOT NULL'],
        ['counterparty', 'VARCHAR(255)', '对方单位', '-'],
        ['counterparties', 'TEXT', '多方JSON', '-'],
        ['amount', 'DECIMAL(18,2)', '金额', '-'],
        ['start_date', 'DATE', '开始日期', '-'],
        ['end_date', 'DATE', '结束日期', '-'],
        ['status', 'VARCHAR(20)', '状态', 'DEFAULT \'DRAFT\''],
        ['content', 'TEXT', '内容', '-'],
        ['attachment', 'VARCHAR(255)', '旧附件字段', '-'],
        ['attachments', 'TEXT', '附件JSON', '-'],
        ['remark', 'TEXT', '备注', '-'],
        ['creator_id', 'BIGINT', '创建人', '-'],
        ['create_time', 'TIMESTAMP', '创建时间', 'DEFAULT CURRENT_TIMESTAMP'],
        ['update_time', 'TIMESTAMP', '更新时间', 'DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, 'template_variable (模板变量表)', level=3)
    headers = ['字段', '类型', '说明', '约束']
    rows = [
        ['id', 'BIGINT', '主键', 'AUTO_INCREMENT'],
        ['code', 'VARCHAR(100)', '变量编码', 'NOT NULL UNIQUE'],
        ['name', 'VARCHAR(100)', '变量名称', 'NOT NULL'],
        ['label', 'VARCHAR(100)', '显示名称', '-'],
        ['type', 'VARCHAR(20)', '变量类型', 'NOT NULL DEFAULT \'text\''],
        ['category', 'VARCHAR(50)', '所属分类', '-'],
        ['default_value', 'VARCHAR(500)', '默认值', '-'],
        ['description', 'VARCHAR(500)', '变量描述', '-'],
        ['required', 'INT', '是否必填', 'DEFAULT 0'],
        ['sort_order', 'INT', '排序', 'DEFAULT 0'],
        ['status', 'INT', '状态', 'DEFAULT 1'],
        ['created_at', 'TIMESTAMP', '创建时间', 'DEFAULT CURRENT_TIMESTAMP'],
        ['updated_at', 'TIMESTAMP', '更新时间', 'DEFAULT CURRENT_TIMESTAMP'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, 'contract_type_field (合同类型字段配置表)', level=3)
    headers = ['字段', '类型', '说明', '约束']
    rows = [
        ['id', 'BIGINT', '主键', 'AUTO_INCREMENT'],
        ['contract_type', 'VARCHAR(50)', '合同类型', 'NOT NULL'],
        ['field_key', 'VARCHAR(50)', '字段标识', 'NOT NULL'],
        ['field_label', 'VARCHAR(100)', '字段标签', '-'],
        ['field_label_en', 'VARCHAR(100)', '字段标签(英)', '-'],
        ['field_type', 'VARCHAR(20)', '字段类型', 'NOT NULL DEFAULT \'text\''],
        ['required', 'BOOLEAN', '是否必填', 'DEFAULT FALSE'],
        ['show_in_list', 'BOOLEAN', '列表显示', 'DEFAULT TRUE'],
        ['show_in_form', 'BOOLEAN', '表单显示', 'DEFAULT TRUE'],
        ['field_order', 'INT', '字段顺序', 'DEFAULT 0'],
        ['placeholder', 'VARCHAR(200)', '占位提示', '-'],
        ['placeholder_en', 'VARCHAR(200)', '占位提示(英)', '-'],
        ['default_value', 'VARCHAR(200)', '默认值', '-'],
        ['options', 'TEXT', '选项配置(JSON)', '-'],
        ['min_value', 'DECIMAL(18,2)', '最小值', '-'],
        ['max_value', 'DECIMAL(18,2)', '最大值', '-'],
        ['created_at', 'TIMESTAMP', '创建时间', 'DEFAULT CURRENT_TIMESTAMP'],
        ['updated_at', 'TIMESTAMP', '更新时间', 'DEFAULT CURRENT_TIMESTAMP'],
        ['UNIQUE KEY', '-', '唯一约束', '(contract_type, field_key)'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # 6. API接口规范
    add_heading_custom(doc, '6. API接口规范', level=1)
    
    add_heading_custom(doc, '6.1 响应格式', level=2)
    add_paragraph_custom(doc, '{\n  "code": 200,\n  "message": "success",\n  "data": {}\n}')
    
    add_heading_custom(doc, '6.2 主要接口', level=2)
    
    add_heading_custom(doc, '认证接口', level=3)
    headers = ['方法', '路径', '说明']
    rows = [
        ['POST', '/api/auth/login', '登录'],
        ['POST', '/api/auth/register', '注册'],
        ['POST', '/api/auth/refresh', '刷新Token'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '合同接口', level=3)
    headers = ['方法', '路径', '说明']
    rows = [
        ['GET', '/api/contracts', '合同列表'],
        ['GET', '/api/contracts/{id}', '合同详情'],
        ['POST', '/api/contracts', '创建合同'],
        ['PUT', '/api/contracts/{id}', '更新合同'],
        ['DELETE', '/api/contracts/{id}', '删除合同'],
        ['POST', '/api/contracts/{id}/submit', '提交审批'],
        ['POST', '/api/contracts/{id}/approve', '审批通过'],
        ['POST', '/api/contracts/{id}/reject', '审批拒绝'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # 7. 安全特性
    add_heading_custom(doc, '7. 安全特性', level=1)
    
    add_heading_custom(doc, '7.1 认证与授权', level=2)
    headers = ['特性', '说明', '状态']
    rows = [
        ['JWT Token认证', '无状态认证机制', '✅ 已完成'],
        ['Token自动刷新', '无感续期机制', '✅ 已完成'],
        ['BCrypt密码加密', '安全密码存储', '✅ 已完成'],
        ['RBAC权限控制', '基于角色的访问控制', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '7.2 API安全', level=2)
    headers = ['特性', '说明', '状态']
    rows = [
        ['速率限制', '60次/分钟防护', '✅ 已完成'],
        ['CORS配置', '跨域安全配置', '✅ 已完成'],
        ['文件路径防护', '防止路径遍历攻击', '✅ 已完成'],
        ['文件类型验证', '白名单验证', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '7.3 审计日志', level=2)
    headers = ['特性', '说明', '状态']
    rows = [
        ['AOP自动记录', '切面编程记录操作', '✅ 已完成'],
        ['操作追溯', '用户、时间、动作、模块', '✅ 已完成'],
    ]
    create_feature_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # 8. 部署说明
    add_heading_custom(doc, '8. 部署说明', level=1)
    
    add_heading_custom(doc, '8.1 部署架构', level=2)
    add_paragraph_custom(doc, '┌─────────────────┐')
    add_paragraph_custom(doc, '│   Nginx (前端)   │  ← 端口 3000')
    add_paragraph_custom(doc, '│   Vue 3 SPA     │')
    add_paragraph_custom(doc, '└────────┬────────┘')
    add_paragraph_custom(doc, '         │')
    add_paragraph_custom(doc, '         ▼')
    add_paragraph_custom(doc, '┌─────────────────┐')
    add_paragraph_custom(doc, '│  Spring Boot    │  ← 端口 8081')
    add_paragraph_custom(doc, '│  REST API       │')
    add_paragraph_custom(doc, '└────────┬────────┘')
    add_paragraph_custom(doc, '         │')
    add_paragraph_custom(doc, '         ▼')
    add_paragraph_custom(doc, '┌─────────────────┐')
    add_paragraph_custom(doc, '│  H2 / MySQL     │  ← 数据库')
    add_paragraph_custom(doc, '└─────────────────┘')
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '8.2 环境变量', level=2)
    add_paragraph_custom(doc, 'export JWT_SECRET="your-secure-secret-key"')
    add_paragraph_custom(doc, 'export CORS_ORIGINS="http://localhost:3000"')
    
    add_heading_custom(doc, '8.3 启动命令', level=2)
    add_paragraph_custom(doc, '# 后端')
    add_paragraph_custom(doc, 'cd backend && mvn spring-boot:run')
    add_paragraph_custom(doc, '')
    add_paragraph_custom(doc, '# 前端')
    add_paragraph_custom(doc, 'cd frontend && npm run dev')
    
    # 保存文档
    doc.save('Toy_Contract_项目功能说明书.docx')
    print("✅ 文档生成成功: Toy_Contract_项目功能说明书.docx")

if __name__ == '__main__':
    generate_document()
